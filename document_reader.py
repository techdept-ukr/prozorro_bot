import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_bytes(content: bytes, filename: str, mime_type: str = "") -> str:
    """Extract text from document bytes based on format."""
    fname = filename.lower()
    mime = mime_type.lower()

    try:
        if fname.endswith(".pdf") or "pdf" in mime:
            return _extract_pdf(content, filename)
        elif fname.endswith(".docx") or "wordprocessingml" in mime or "docx" in mime:
            return _extract_docx(content, filename)
        elif fname.endswith(".doc") or "msword" in mime:
            return _extract_doc(content, filename)
        elif fname.endswith(".xlsx") or "spreadsheetml" in mime:
            return _extract_xlsx(content, filename)
        elif fname.endswith(".txt") or "text/plain" in mime:
            return content.decode("utf-8", errors="replace")
        elif fname.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif")):
            return _extract_image_ocr(content, filename)
        else:
            # Try as text
            try:
                return content.decode("utf-8", errors="replace")
            except Exception:
                return f"[Не вдалося прочитати файл: {filename}]"
    except Exception as e:
        logger.warning(f"Error extracting text from {filename}: {e}")
        return f"[Помилка читання файлу {filename}: {e}]"


def _extract_pdf(content: bytes, filename: str) -> str:
    """Extract text from PDF, fallback to OCR for scanned pages."""
    import pdfplumber

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if len(page_text.strip()) < 50:
                    # Likely scanned — use OCR
                    ocr_text = _ocr_pdf_page(page)
                    text_parts.append(ocr_text if ocr_text else f"[Сторінка {i+1}: скан без тексту]")
                else:
                    text_parts.append(page_text)
    except Exception as e:
        logger.warning(f"pdfplumber failed for {filename}: {e}")
        # Full OCR fallback
        return _full_pdf_ocr(content, filename)

    return "\n\n".join(text_parts)


def _ocr_pdf_page(page) -> str:
    """OCR a single pdfplumber page using pytesseract."""
    try:
        import pytesseract
        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img, lang="ukr+rus+eng")
    except Exception as e:
        logger.warning(f"OCR page failed: {e}")
        return ""


def _full_pdf_ocr(content: bytes, filename: str) -> str:
    """Full OCR for scanned PDFs using pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(content, dpi=200)
        texts = []
        for img in images:
            t = pytesseract.image_to_string(img, lang="ukr+rus+eng")
            texts.append(t)
        return "\n\n".join(texts)
    except Exception as e:
        logger.warning(f"Full PDF OCR failed for {filename}: {e}")
        return f"[Не вдалося розпізнати PDF: {filename}]"


def _extract_docx(content: bytes, filename: str) -> str:
    """Extract text from .docx file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"DOCX extraction failed for {filename}: {e}")
        return f"[Помилка читання DOCX: {filename}]"


def _extract_doc(content: bytes, filename: str) -> str:
    """Extract text from .doc (old Word) file using antiword or python-docx2txt."""
    try:
        import docx2txt
        return docx2txt.process(io.BytesIO(content))
    except Exception:
        try:
            import subprocess, tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
                f.write(content)
                tmp_path = f.name
            result = subprocess.run(
                ["antiword", tmp_path], capture_output=True, text=True, timeout=30
            )
            os.unlink(tmp_path)
            return result.stdout or f"[Порожній DOC: {filename}]"
        except Exception as e:
            return f"[Не вдалося прочитати DOC: {filename} — {e}]"


def _extract_xlsx(content: bytes, filename: str) -> str:
    """Extract text from .xlsx spreadsheet."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"=== Аркуш: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(v) for v in row if v is not None)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Помилка читання XLSX: {filename} — {e}]"


def _extract_image_ocr(content: bytes, filename: str) -> str:
    """OCR an image file."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        return pytesseract.image_to_string(img, lang="ukr+rus+eng")
    except Exception as e:
        return f"[Не вдалося розпізнати зображення: {filename} — {e}]"
