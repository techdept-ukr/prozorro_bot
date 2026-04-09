import os
import re
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def generate_docx_report(data: dict) -> str:
    """Generate a .docx analysis report. Returns path to the file."""
    doc = Document()
    _setup_styles(doc)

    tender = data["tender"]
    tender_id = data["tender_id"]

    # ── Title Page ──────────────────────────────────────────────────
    _add_title_block(doc, tender, tender_id, data["generated_at"])

    # ── Section 1: Customer Analysis ────────────────────────────────
    _add_section_heading(doc, "1. АНАЛІЗ ЗАМОВНИКА ТА ТЕХНІЧНИХ ВИМОГ", level=1)
    _add_analysis_text(doc, data["customer_analysis"])
    doc.add_page_break()

    # ── Section 2: Procurement Analysis ─────────────────────────────
    _add_section_heading(doc, "2. АНАЛІЗ ЗАКУПІВЛІ", level=1)
    _add_analysis_text(doc, data["procurement_analysis"])
    doc.add_page_break()

    # ── Section 3: Participants Analysis ────────────────────────────
    _add_section_heading(doc, "3. АНАЛІЗ УЧАСНИКІВ", level=1)
    _add_analysis_text(doc, data["participants_analysis"])
    doc.add_page_break()

    # ── Section 4: Raw Tender Info ───────────────────────────────────
    _add_section_heading(doc, "4. ЗАГАЛЬНІ ВІДОМОСТІ ПРО ТЕНДЕР", level=1)
    _add_tender_metadata_table(doc, tender)

    # ── Footer ───────────────────────────────────────────────────────
    _add_footer(doc, tender_id, data["generated_at"])

    # Save
    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False,
        prefix=f"tender_{tender_id}_"
    )
    doc.save(tmp.name)
    return tmp.name


def _setup_styles(doc: Document):
    """Configure default document styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def _add_title_block(doc: Document, tender: dict, tender_id: str, generated_at: str):
    # Header bar
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("АНАЛІЗ ТЕНДЕРНОЇ ЗАКУПІВЛІ")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    # Tender title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_para.add_run(tender.get("title", "Без назви"))
    run2.bold = True
    run2.font.size = Pt(14)

    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    _set_table_col_width(table, 0, Cm(5))
    _set_table_col_width(table, 1, Cm(11))

    rows_data = [
        ("ID тендера:", tender_id),
        ("Статус:", _status_ua(tender.get("status", ""))),
        ("Очікувана вартість:", f"{tender.get('value', {}).get('amount', 'н/д')} {tender.get('value', {}).get('currency', 'UAH')}"),
        ("Процедура:", tender.get("procurementMethodType", "н/д")),
        ("Дата публікації:", tender.get("datePublished", "н/д")[:10] if tender.get("datePublished") else "н/д"),
        ("Дата аналізу:", generated_at),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value

    doc.add_page_break()


def _add_section_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    para.space_before = Pt(12)
    para.space_after = Pt(6)
    _add_horizontal_rule(doc)


def _add_analysis_text(doc: Document, text: str):
    """Parse markdown-like AI output into formatted Word paragraphs."""
    if not text:
        doc.add_paragraph("Аналіз недоступний.")
        return

    lines = text.split("\n")
    for line in lines:
        line = line.rstrip()

        if not line:
            doc.add_paragraph()
            continue

        # Heading 2: ## or **TEXT** alone on line
        if re.match(r'^#{1,2}\s+', line):
            clean = re.sub(r'^#+\s+', '', line)
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            _add_sub_heading(doc, clean)

        # Heading 3: ### or numbered like "1. " at start
        elif re.match(r'^###\s+', line):
            clean = re.sub(r'^###\s+', '', line)
            _add_sub_heading(doc, clean, size=Pt(12))

        # Bullet points
        elif re.match(r'^[-•*]\s+', line):
            clean = re.sub(r'^[-•*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, clean)

        # Numbered list
        elif re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            clean = re.sub(r'^\d+\.\s+', '', line)
            _add_formatted_run(p, clean)

        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, line)


def _add_formatted_run(para, text: str):
    """Add text with **bold** and _italic_ markdown parsed inline."""
    parts = re.split(r'(\*\*[^*]+\*\*|_[^_]+_)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _add_sub_heading(doc: Document, text: str, size=Pt(13)):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    para.space_before = Pt(8)


def _add_tender_metadata_table(doc: Document, tender: dict):
    """Add raw tender metadata as a readable table."""
    items = tender.get("items", [])
    if items:
        _add_sub_heading(doc, "Предмет закупівлі")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Опис", "CPV", "Кількість", "Одиниця"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for item in items[:30]:
            row = table.add_row().cells
            row[0].text = item.get("description", "")[:100]
            row[1].text = item.get("classification", {}).get("id", "")
            row[2].text = str(item.get("quantity", ""))
            row[3].text = item.get("unit", {}).get("name", "")

    bids = tender.get("bids", [])
    if bids:
        doc.add_paragraph()
        _add_sub_heading(doc, "Учасники та цінові пропозиції")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Учасник", "ЄДРПОУ", "Ціна (UAH)", "Статус"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for bid in bids:
            tenderer = bid.get("tenderers", [{}])[0]
            row = table.add_row().cells
            row[0].text = tenderer.get("name", "")[:60]
            row[1].text = tenderer.get("identifier", {}).get("id", "")
            row[2].text = str(bid.get("value", {}).get("amount", "н/д"))
            row[3].text = _status_ua(bid.get("status", ""))


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_footer(doc: Document, tender_id: str, generated_at: str):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Аналіз тендера {tender_id} | Згенеровано: {generated_at} | ProZorro Analyzer Bot"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _set_table_col_width(table, col_idx, width):
    for row in table.rows:
        row.cells[col_idx].width = width


def _status_ua(status: str) -> str:
    mapping = {
        "active.tendering": "Прийом пропозицій",
        "active.enquiries": "Уточнення",
        "active.pre-qualification": "Прекваліфікація",
        "active.qualification": "Кваліфікація",
        "active.awarded": "Визначено переможця",
        "complete": "Завершено",
        "cancelled": "Скасовано",
        "unsuccessful": "Не відбулось",
        "active": "Активний",
        "pending": "Очікує",
        "invalid": "Відхилено",
    }
    return mapping.get(status, status)
